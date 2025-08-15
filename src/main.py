#!/usr/bin/env python3
"""
Main script for the Performance Review Tracking System.

This script provides end-to-end automation for generating performance review reports,
including data processing, LLM analysis, and report generation.
"""

import argparse
import json
import os
import sys
import subprocess
from datetime import datetime
from typing import Optional, Union, Dict
import pandas as pd
from docx import Document
import markdown
from markdown.extensions import fenced_code, tables
import site

# Import LLM analyzer module
try:
    from .llm_analyzer import run_llm_analysis, validate_llm_config
except ImportError:
    try:
        from llm_analyzer import run_llm_analysis, validate_llm_config
    except ImportError:
        print("Warning: LLM analyzer module not found. Will use fallback manual analysis.", file=sys.stderr)
        def run_llm_analysis(data_file, review_type, config):
            # Fallback implementation
            from .llm_analyzer import generate_manual_analysis
            return generate_manual_analysis(data_file, review_type, None, config)
        def validate_llm_config(config):
            return []

print(f"Python executable: {sys.executable}")
print(f"Python version: {sys.version}")
print(f"Site packages: {site.getsitepackages()}")

# Import configuration validation module
try:
    from config_validation import load_and_validate_config, ConfigValidationError
except ImportError:
    print("Warning: config_validation module not found. Configuration management disabled.", file=sys.stderr)
    def load_and_validate_config(*args, **kwargs):
        return {}
    class ConfigValidationError(Exception):
        pass

# Data processor functions
def load_data(file_path: str) -> pd.DataFrame:
    """
    Load data from CSV or Excel file.
    """
    print("DEBUG: load_data - START", file_path) # DEBUG LOG
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    file_ext = os.path.splitext(file_path)[1].lower()
    try:
        if file_ext == '.csv':
            print("DEBUG: load_data - Loading CSV", file_path) # DEBUG LOG
            df = pd.read_csv(file_path)
            print("DEBUG: load_data - CSV loaded successfully", file_path) # DEBUG LOG
            return df
        elif file_ext in ['.xlsx', '.xls']:
            print("DEBUG: load_data - Loading Excel", file_path) # DEBUG LOG
            df = pd.read_excel(file_path)
            print("DEBUG: load_data - Excel loaded successfully", file_path) # DEBUG LOG
            return df
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    except Exception as e:
        print(f"DEBUG: load_data - ERROR loading data: {str(e)}") # DEBUG LOG
        raise ValueError(f"Error loading data from {file_path}: {str(e)}")

    print("DEBUG: load_data - END") # DEBUG LOG










def filter_by_date_range(
    data: pd.DataFrame,
    start_date: Union[str, datetime],
    end_date: Union[str, datetime]
) -> pd.DataFrame:
    """
    Filter entries by date range for Annual Review.
    """
    print("DEBUG: filter_by_date_range - START", start_date, end_date) # DEBUG LOG
    if 'Date' not in data.columns:
        raise ValueError("DataFrame must contain a 'Date' column")

    # Make a copy to avoid modifying the original
    filtered_data = data.copy()

    # Try to convert dates to datetime, with error handling
    try:
        filtered_data['Date'] = pd.to_datetime(filtered_data['Date'], errors='raise')
    except Exception as e:
        raise ValueError(f"Invalid date format in data: {str(e)}")

    # Convert string dates to datetime if needed
    if isinstance(start_date, str):
        try:
            start_date = datetime.strptime(start_date, '%Y-%m-%d')
        except ValueError:
            raise ValueError(f"Invalid start date format: {start_date}. Use YYYY-MM-DD")

    if isinstance(end_date, str):
        try:
            end_date = datetime.strptime(end_date, '%Y-%m-%d')
        except ValueError:
            raise ValueError(f"Invalid end date format: {end_date}. Use YYYY-MM-DD")

    start_ts = pd.Timestamp(start_date)
    end_ts = pd.Timestamp(end_date)

    # Filter by date range
    mask = filtered_data['Date'].between(start_ts, end_ts)
    print("DEBUG: filter_by_date_range - END") # DEBUG LOG
    return filtered_data[mask].reset_index(drop=True)


def normalize_data_structure(data: pd.DataFrame) -> pd.DataFrame:
    """
    Normalize data structure by adding missing columns with default values.
    """
    print("DEBUG: normalize_data_structure - START") # DEBUG LOG
    
    # Create a copy to avoid modifying original
    normalized_data = data.copy()
    
    # Required columns and their defaults
    required_columns_defaults = {
        'Date': None,  # Must be provided
        'Title': None,  # Must be provided  
        'Description': '',
        'Acceptance Criteria': '',
        'Success Notes': '',
        'Impact': 'Medium',
        'Self Rating': 2,
        'Rating Justification': 'Performance meets expectations'
    }
    
    # Add missing columns with defaults
    for col, default_value in required_columns_defaults.items():
        if col not in normalized_data.columns:
            if default_value is not None:
                normalized_data[col] = default_value
                print(f"DEBUG: Added missing column '{col}' with default value: {default_value}")
            else:
                # This is a critical error - Date and Title must be provided
                raise ValueError(f"Missing critical column: {col}")
    
    # Clean up any NaN values in non-critical columns
    normalized_data['Description'] = normalized_data['Description'].fillna('')
    normalized_data['Acceptance Criteria'] = normalized_data['Acceptance Criteria'].fillna('')
    normalized_data['Success Notes'] = normalized_data['Success Notes'].fillna('')
    normalized_data['Self Rating'] = normalized_data['Self Rating'].fillna(2)
    normalized_data['Rating Justification'] = normalized_data['Rating Justification'].fillna('Performance meets expectations')
    
    print("DEBUG: normalize_data_structure - END") # DEBUG LOG
    return normalized_data


def validate_data(data: pd.DataFrame) -> None:
    """
    Validate the structure and content of the input data.
    """
    print("DEBUG: validate_data - START") # DEBUG LOG
    
    # Critical columns that must exist
    critical_columns = ['Date', 'Title']
    missing_critical = [col for col in critical_columns if col not in data.columns]
    if missing_critical:
        raise ValueError(f"Missing critical columns: {', '.join(missing_critical)}")

    # Check for empty titles (after normalization, descriptions can be empty and filled)
    empty_titles = data['Title'].isnull() | (data['Title'].astype(str).str.strip() == '')
    if empty_titles.any():
        raise ValueError("Found entries with empty titles")

    # Validate Impact values if present
    if 'Impact' in data.columns:
        valid_impacts = ['High', 'Medium', 'Low']
        invalid_impacts = data[~data['Impact'].isin(valid_impacts)]['Impact'].unique()
        if len(invalid_impacts) > 0:
            raise ValueError(
                f"Invalid Impact values found: {', '.join(map(str, invalid_impacts))}. "
                f"Valid values are: {', '.join(valid_impacts)}"
            )

    # Validate Self Rating values if present (1-3)
    if 'Self Rating' in data.columns:
        valid_ratings = [1, 2, 3]
        # Convert ratings to numeric, replacing any non-numeric values with NaN
        data['Self Rating'] = pd.to_numeric(data['Self Rating'], errors='coerce').fillna(2)
        invalid_ratings = data[~data['Self Rating'].isin(valid_ratings)]['Self Rating'].dropna().unique()
        if len(invalid_ratings) > 0:
            raise ValueError(
                f"Invalid Self Rating values found: {', '.join(map(str, invalid_ratings))}. "
                "Valid values are: 1 (Poor), 2 (Good), 3 (Excellent)"
            )

    print("DEBUG: validate_data - END") # DEBUG LOG


def prepare_data_for_analysis(
    data: pd.DataFrame,
    review_type: str,
    year: Optional[str] = None,
    output_dir: str = 'data'
) -> str:
    """
    Prepare data for analysis.
    """
    print("DEBUG: prepare_data_for_analysis - START", review_type, year, output_dir) # DEBUG LOG
    # Validate review type
    if review_type.lower() not in ['annual', 'competency']:
        raise ValueError(
            "Invalid review type. Must be either 'annual' or 'competency'"
        )

    # Check if data is empty (after header)
    if len(data) == 0:
        raise ValueError("No data found in the input file")

    # Normalize data structure (add missing columns with defaults)
    data = normalize_data_structure(data)
    
    # Validate data
    validate_data(data)

    # Always try to convert 'Date' column to datetime for consistency
    try:
        data['Date'] = pd.to_datetime(data['Date'], errors='raise')
    except Exception as e:
        raise ValueError(f"Invalid date format in data: {str(e)}")

    if review_type.lower() == 'annual':
        if not year:
            raise ValueError("Year is required for annual reviews")

        # Annual review covers Sept to Aug
        start_date = f"{int(year)-1}-09-01"
        end_date = f"{year}-08-31"
        data = filter_by_date_range(data, start_date, end_date)

    # Ensure the output directory exists
    os.makedirs(output_dir, exist_ok=True)

    # Convert to JSON format for analysis
    output_path = os.path.join(output_dir, f"processed_{review_type.lower()}.json")

    # Convert dates to ISO format strings for JSON serialization
    # First ensure we have datetime type in the Date column
    if hasattr(data['Date'], 'dt'):
        data['Date'] = data['Date'].dt.strftime('%Y-%m-%d')
    else:
        # If somehow we don't have datetime objects, try a safe conversion
        data['Date'] = data['Date'].apply(lambda x: x.strftime('%Y-%m-%d') if hasattr(x, 'strftime') else str(x))

    # Save to file for analysis access
    print(f"DEBUG: prepare_data_for_analysis - Saving processed data to: {output_path}") # DEBUG LOG
    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(json.loads(data.to_json(orient='records')), f, indent=2)
    except Exception as e:
        print(f"DEBUG: prepare_data_for_analysis - ERROR saving JSON: {str(e)}") # DEBUG LOG
        raise
    print(f"DEBUG: prepare_data_for_analysis - Saved processed data to: {output_path}") # DEBUG LOG
    print("DEBUG: prepare_data_for_analysis - END", output_path) # DEBUG LOG
    return output_path

# Report generator functions
def load_template(review_type):
    """
    Load the appropriate review template.
    """
    print("DEBUG: load_template - START", review_type) # DEBUG LOG
    if review_type == "competency":
        template_path = f"templates/competency_assessment_template.md"
    else:
        template_path = f"templates/annual_review_template.md"
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template file not found: {template_path}")

    with open(template_path, 'r', encoding='utf-8') as f:
        template_content = f.read()
    print("DEBUG: load_template - END", template_path) # DEBUG LOG
    return template_content

def load_analysis(file_path):
    """
    Load the analysis generated by the analyzer.
    """
    print("DEBUG: load_analysis - START", file_path) # DEBUG LOG
    analysis_path = file_path.replace('processed_', 'analyzed_').replace('.json', '.md')

    if not os.path.exists(analysis_path):
        raise FileNotFoundError(
            f"Analysis file not found: {analysis_path}\n\n"
            "Please complete these steps first:\n"
            "1. Open VS Code\n"
            "2. Switch to Performance Analyst mode\n"
            f"3. Use the command: @{file_path} Please analyze this data and generate a structured report\n"
            "4. Save the analysis to: " + analysis_path
        )

    with open(analysis_path, 'r', encoding='utf-8') as f:
        analyzed_content = f.read()
    print("DEBUG: load_analysis - Content read from analysis file:\n", analyzed_content) # DEBUG LOG
    print("DEBUG: load_analysis - END", analysis_path) # DEBUG LOG
    return analyzed_content

def markdown_to_docx(markdown_text, output_path):
    """
    Convert markdown formatted text to a DOCX document.
    """
    print("DEBUG: markdown_to_docx - START", output_path) # DEBUG LOG
    if not markdown_text.strip():
        raise ValueError("Markdown text cannot be empty")

    # Create new Document
    doc = Document()

    # Process markdown line by line
    in_list = False
    lines = markdown_text.split('\n')

    for line in lines:
        line = line.rstrip()
        if not line:
            # Empty line
            continue

        # Check for headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        # Check for list items
        elif line.startswith('- '):
            if not in_list:
                in_list = True
            doc.add_paragraph(line[2:], style='List Bullet')
        # Regular paragraph
        else:
            in_list = False
            doc.add_paragraph(line)

    # Save the document
    print(f"DEBUG: markdown_to_docx - Saving DOCX to: {output_path}") # DEBUG LOG
    doc.save(output_path)
    print("DEBUG: markdown_to_docx - END", output_path) # DEBUG LOG
    return output_path

def generate_final_report(data_file, review_type, output_format='markdown', output_path=None):
    """
    Generate final formatted report by combining analysis with template.
    """
    print("DEBUG: generate_final_report - START", data_file, review_type, output_format, output_path) # DEBUG LOG
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    current_year = datetime.now().year

    # If no output path is specified, create a default one
    if not output_path:
        output_path = f"output/{review_type.lower()}_review_{timestamp}.{'md' if output_format.lower() == 'markdown' else 'docx'}"

    # Ensure the output directory exists
    output_dir = os.path.dirname(output_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    try:
        # Load analysis
        analyzed_content = load_analysis(data_file)
        print("DEBUG: generate_final_report - Loaded analysis content:\n", analyzed_content) # DEBUG LOG

        # Load and apply the template
        template = load_template(review_type)

        # For annual review, modify the year and date range
        if review_type.lower() == 'annual':
            # Replace [Year] with the current year
            template = template.replace('[Year]', str(current_year))
            # Replace date range
            template = template.replace('[September 1, YYYY - August 31, YYYY]',
                                    f'[September 1, {current_year-1} - August 31, {current_year}]')

        # Handle empty analysis content
        if not analyzed_content.strip():
            analyzed_content = "No accomplishments found in the specified date range."

        # Check if the analyzed content is LLM-generated (contains structured markdown)
        if analyzed_content.startswith('#') and ('## Executive Summary' in analyzed_content or '## Criterion Analysis' in analyzed_content):
            # LLM-generated content is already properly formatted - use directly
            report_content = analyzed_content
            print("DEBUG: Using LLM-generated content directly (already formatted)")
        else:
            # Traditional analysis - use template formatting
            # For now, just use the analyzed content since templates don't have proper placeholders
            report_content = analyzed_content
            print("DEBUG: Using analyzed content directly (template formatting needs {analyzed_content} placeholder)")
            
        # For annual review, add current year to the title if it's LLM content
        if review_type.lower() == 'annual' and report_content.startswith('# Annual Review'):
            report_content = report_content.replace('# Annual Review', f'# Annual Review {current_year}')
        print("DEBUG: generate_final_report - Report content after template formatting:\n", report_content) # DEBUG LOG

        if output_format.lower() == 'markdown':
            print(f"DEBUG: generate_final_report - Saving markdown report to: {output_path}") # DEBUG LOG
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(report_content)
            print(f"DEBUG: generate_final_report - Saved markdown report to: {output_path}") # DEBUG LOG
        elif output_format.lower() == 'docx':
            # If output path is for a docx, but we need a temporary md file
            if output_path.endswith('.docx'):
                md_path = output_path.replace('.docx', '.md')
            else:
                md_path = f"{output_path}.md"

            # Save markdown version
            print(f"DEBUG: generate_final_report - Saving temporary markdown for DOCX conversion to: {md_path}") # DEBUG LOG
            with open(md_path, 'w', encoding='utf-8') as f:
                f.write(report_content)
            print(f"DEBUG: generate_final_report - Saved temporary markdown for DOCX conversion to: {md_path}") # DEBUG LOG

            # Convert to DOCX
            docx_path = output_path if output_path.endswith('.docx') else f"{output_path}.docx"
            markdown_to_docx(report_content, docx_path)
            output_path = docx_path

            # Clean up temporary markdown file if it was just for conversion
            if md_path != output_path and os.path.exists(md_path):
                os.remove(md_path)

    except NotImplementedError as e:
        # Pass through the usage instructions
        raise NotImplementedError(str(e))
    except Exception as e:
        print(f"DEBUG: generate_final_report - ERROR: {str(e)}") # DEBUG LOG
        raise ValueError(f"Error generating report: {str(e)}")

    print("DEBUG: generate_final_report - END", output_path) # DEBUG LOG
    return output_path




def generate_manual_analysis(data_file: str, review_type: str) -> str:
    """
    Generate manual analysis when AI analysis fails or is empty.
    """
    import json
    
    # Load the processed data
    with open(data_file, 'r') as f:
        data = json.load(f)
    
    total_accomplishments = len(data)
    if total_accomplishments == 0:
        return "No accomplishments found in the processed data."
    
    # Analyze impact distribution
    high_impact = len([item for item in data if item.get('Impact') == 'High'])
    medium_impact = len([item for item in data if item.get('Impact') == 'Medium'])
    low_impact = len([item for item in data if item.get('Impact') == 'Low'])
    
    # Extract date range
    dates = [item['Date'] for item in data if 'Date' in item]
    date_range = f"{min(dates)} to {max(dates)}" if dates else "Date range not available"
    
    # Analyze key themes from titles
    titles = [item.get('Title', '').lower() for item in data]
    
    if review_type.lower() == 'annual':
        return f"""# Annual Review Analysis

## Executive Summary
Successfully completed **{total_accomplishments} work items** during the review period ({date_range}), demonstrating consistent delivery and professional excellence across multiple domains.

**Impact Distribution:**
- High Impact: {high_impact} items ({high_impact/total_accomplishments*100:.1f}%)
- Medium Impact: {medium_impact} items ({medium_impact/total_accomplishments*100:.1f}%)
- Low Impact: {low_impact} items ({low_impact/total_accomplishments*100:.1f}%)

## Criterion Analysis

### 1. Communication
**Self Rating: 2 (Good)**

**How I Met This Criterion:**
Based on {total_accomplishments} completed work items, demonstrated clear communication through detailed technical documentation, comprehensive acceptance criteria, and thorough project specifications. All work items include clear descriptions and success criteria, indicating effective communication of requirements and deliverables.

### 2. Flexibility
**Self Rating: 2 (Good)**

**How I Met This Criterion:**
Successfully adapted to diverse project requirements and changing priorities across {total_accomplishments} work items. Maintained effectiveness while working on various technical domains and adjusting to evolving project needs throughout the review period.

### 3. Initiative
**Self Rating: 2 (Good)**

**How I Met This Criterion:**
Proactively completed {total_accomplishments} work items with {high_impact} classified as high impact, demonstrating self-direction and proactive problem-solving. Consistently identified and addressed project requirements before they became critical issues.

### 4. Member Service
**Self Rating: 2 (Good)**

**How I Met This Criterion:**
Delivered consistent service to internal stakeholders through reliable completion of {total_accomplishments} work items. Focused on meeting project requirements and supporting organizational objectives through quality deliverables.

### 5. Personal Credibility
**Self Rating: 3 (Excellent)**

**How I Met This Criterion:**
Demonstrated exceptional reliability with 100% completion rate across all {total_accomplishments} assigned work items. Consistently met commitments and followed through on all project responsibilities as evidenced by comprehensive success documentation.

### 6. Quality and Quantity of Work
**Self Rating: 3 (Excellent)**

**How I Met This Criterion:**
Achieved outstanding productivity with {total_accomplishments} completed work items while maintaining high quality standards. {high_impact} high-impact deliverables demonstrate significant value contribution to organizational objectives.

### 7. Teamwork
**Self Rating: 2 (Good)**

**How I Met This Criterion:**
Collaborated effectively on {total_accomplishments} work items, supporting team goals through consistent delivery and knowledge sharing. Contributed to collective success through reliable execution of assigned responsibilities.

## Overall Summary

Delivered exceptional performance with {total_accomplishments} successfully completed work items spanning {date_range}. Demonstrated strong professional competencies across all evaluation criteria, with particular excellence in personal credibility and work quality/quantity.

**Key Accomplishments:**
- {total_accomplishments} work items completed successfully
- {high_impact} high-impact deliverables providing significant organizational value
- Consistent delivery maintaining quality standards throughout the review period
- 100% completion rate demonstrating reliability and commitment

**Areas of Strength:**
- Exceptional reliability and follow-through on commitments
- High-volume delivery without compromising quality
- Effective adaptation to diverse project requirements
- Strong technical execution and problem-solving capabilities

_Analysis generated from {total_accomplishments} completed work items during the period {date_range}_
"""

    else:  # competency review
        return f"""# Competency Assessment Analysis

## Executive Summary
Successfully completed **{total_accomplishments} work items** during the assessment period ({date_range}), demonstrating technical competency and professional development across multiple technical domains.

**Impact Distribution:**
- High Impact: {high_impact} items
- Medium Impact: {medium_impact} items  
- Low Impact: {low_impact} items

## Competency Analysis

Based on {total_accomplishments} completed work items, this assessment evaluates performance across 13 professional competency areas:

### Programming/Software Development
**Rating: 3 (Practicing)**
Demonstrated consistent application of programming and development skills across multiple work items with independent execution and quality deliverables.

### Solution Architecture  
**Rating: 2 (Developing)**
Applied architectural thinking to technical solutions with growing proficiency in system design and integration approaches.

### Systems Design
**Rating: 2 (Developing)**
Contributed to system design initiatives with increasing understanding of scalability and integration requirements.

### Project Management
**Rating: 3 (Practicing)**
Successfully managed {total_accomplishments} work items from initiation through completion, demonstrating effective project execution skills.

### Requirements Definition
**Rating: 3 (Practicing)**
Consistently delivered work items meeting defined acceptance criteria, showing strong requirements analysis and implementation capabilities.

### Testing
**Rating: 2 (Developing)**
Applied testing practices and validation approaches with growing proficiency in quality assurance methodologies.

### Problem Management
**Rating: 3 (Practicing)**
Effectively resolved technical challenges across {total_accomplishments} work items, demonstrating systematic problem-solving abilities.

### Innovation
**Rating: 2 (Developing)**
Contributed creative solutions and process improvements with increasing confidence in innovative approaches.

### Release/Deployment
**Rating: 2 (Developing)**
Participated in deployment activities with growing understanding of release management and production considerations.

### Accountability
**Rating: 4 (Mastering)**
Demonstrated exceptional accountability with 100% completion rate across all assigned work items and consistent follow-through on commitments.

### Influence
**Rating: 2 (Developing)**
Building influence through reliable delivery and technical contributions, with opportunities for expanded leadership impact.

### Agility
**Rating: 3 (Practicing)**
Successfully adapted to changing requirements and priorities across diverse work items, maintaining effectiveness during transitions.

### Inclusion
**Rating: 2 (Developing)**
Contributed to team success through collaborative execution and knowledge sharing opportunities.

## Overall Assessment

Strong competency development demonstrated through {total_accomplishments} successfully completed work items. Particular strength in accountability, project execution, and technical problem-solving. Opportunities for continued growth in leadership influence and architectural design capabilities.

**Development Recommendations:**
- Continue building solution architecture experience through complex system design projects
- Expand leadership influence through mentoring and cross-team collaboration
- Deepen innovation capabilities through process improvement initiatives

_Assessment based on {total_accomplishments} work items completed during {date_range}_
"""


def generate_review_with_config(
    config: Dict,
    source: str = "csv",
    input_file: Optional[str] = None,
    review_type: str = "competency",
    year: Optional[str] = None,
    output_format: str = 'markdown',
    output_path: Optional[str] = None
) -> str:
    """
    Generate a performance review report from CSV data source.
    
    Args:
        config: Configuration dictionary
        source: Data source (must be "csv")
        input_file: CSV file path (required)
        review_type: Type of review ("annual" or "competency")
        year: Year for annual review
        output_format: Output format ("markdown" or "docx")
        output_path: Custom output path
        
    Returns:
        Path to generated report
    """
    print("DEBUG: generate_review_with_config - START", source, input_file, review_type) # DEBUG LOG
    
    try:
        # Load data from CSV file
        if not input_file:
            raise ValueError("input_file is required for CSV data source")
        if source != "csv":
            raise ValueError("Only CSV data source is supported")
            
        print(f"DEBUG: Loading data from CSV file: {input_file}")
        data = load_data(input_file)
        
        # Continue with existing review generation logic
        return generate_review_from_data(
            data=data,
            review_type=review_type,
            year=year,
            output_format=output_format,
            output_path=output_path,
            config=config
        )
        
    except Exception as e:
        print(f"DEBUG: generate_review_with_config - ERROR: {str(e)}") # DEBUG LOG
        raise




def generate_review_from_data(
    data: pd.DataFrame,
    review_type: str,
    year: Optional[str] = None,
    output_format: str = 'markdown',
    output_path: Optional[str] = None,
    config: Optional[Dict] = None
) -> str:
    """
    Generate a performance review report from DataFrame data.
    
    Args:
        data: DataFrame with performance review data
        review_type: Type of review ("annual" or "competency")  
        year: Year for annual review
        output_format: Output format ("markdown" or "docx")
        output_path: Custom output path
        config: Configuration dictionary (for output directory)
        
    Returns:
        Path to generated report
    """
    print("DEBUG: generate_review_from_data - START", review_type, year, output_format) # DEBUG LOG
    
    try:
        # Prepare data for analysis
        output_dir = "data"
        if config and config.get("processing", {}).get("output_directory"):
            output_dir = config["processing"]["output_directory"]
            
        processed_data_path = prepare_data_for_analysis(
            data,
            review_type,
            year,
            output_dir
        )
        print(f"DEBUG: Data processed and saved to {processed_data_path}")
        
        # Run analysis (LLM integration)
        analysis_path = run_llm_analysis(processed_data_path, review_type, config or {})
        
        # Generate final report
        report_path = generate_final_report(
            processed_data_path,
            review_type,
            output_format,
            output_path
        )
        print(f"DEBUG: Report generated successfully at: {report_path}")
        
        print("DEBUG: generate_review_from_data - END", report_path) # DEBUG LOG
        return report_path
        
    except Exception as e:
        print(f"DEBUG: generate_review_from_data - ERROR: {str(e)}") # DEBUG LOG
        raise


def generate_review(
    input_file: str,
    review_type: str,
    year: Optional[str] = None,
    output_format: str = 'markdown',
    output_path: Optional[str] = None
) -> str:
    """
    Generate a performance review report from start to finish.
    """
    print("DEBUG: generate_review - START", input_file, review_type, year, output_format, output_path) # DEBUG LOG
    try:
        # 1. Load and process data
        print(f"DEBUG: generate_review - Loading data from {input_file}...") # DEBUG LOG
        data = load_data(input_file)
        print(f"DEBUG: generate_review - Data loaded successfully from {input_file}") # DEBUG LOG

        print("DEBUG: generate_review - Processing data...") # DEBUG LOG
        processed_data_path = prepare_data_for_analysis(
            data,
            review_type,
            year,
            'data'  # Output directory
        )
        print(f"DEBUG: generate_review - Data processed and saved to {processed_data_path}") # DEBUG LOG

        # 2. Run LLM analysis (no config available - will use fallback)
        analysis_path = run_llm_analysis(processed_data_path, review_type, {})

        # 3. Generate final report
        print("DEBUG: generate_review - Generating final report...") # DEBUG LOG
        report_path = generate_final_report(
            processed_data_path,
            review_type,
            output_format,
            output_path
        )
        print(f"DEBUG: generate_review - Report generated successfully at: {report_path}") # DEBUG LOG

        print("DEBUG: generate_review - END", report_path) # DEBUG LOG
        return report_path

    except ValueError as e:
        print(f"DEBUG: generate_review - ValueError: {str(e)}") # DEBUG LOG
        print(f"Error: Invalid input - {str(e)}", file=sys.stderr)
        raise
    except RuntimeError as e:
        print(f"DEBUG: generate_review - RuntimeError: {str(e)}") # DEBUG LOG
        print(f"Error: Process failed - {str(e)}", file=sys.stderr)
        raise
    except Exception as e:
        print(f"DEBUG: generate_review - Unexpected Exception: {str(e)}") # DEBUG LOG
        print(f"Error: Unexpected error - {str(e)}", file=sys.stderr)
        raise


def main():
    """Main function to orchestrate the review generation process."""
    parser = argparse.ArgumentParser(
        description="Generate performance review reports from CSV data",
        formatter_class=argparse.RawTextHelpFormatter,
        epilog="""
Examples:
  Generate competency assessment:
    %(prog)s --file data/accomplishments.csv --type competency
  
  Generate annual review:
    %(prog)s --file data/accomplishments.csv --type annual --year 2025
    
  Generate with custom output format:
    %(prog)s --file data/accomplishments.csv --type annual --year 2025 --format docx
        """
    )

    # Configuration arguments
    parser.add_argument(
        "--config",
        default="config.json",
        help="Path to configuration file (default: config.json)"
    )
    parser.add_argument(
        "--create-config",
        action="store_true",
        help="Create example configuration file and exit"
    )
    parser.add_argument(
        "--test-config",
        action="store_true", 
        help="Test configuration and connections, then exit"
    )
    
    # Data arguments
    parser.add_argument(
        "--file",
        required=True,
        help="Path to CSV/Excel data file (required)"
    )
    
    # Review arguments
    parser.add_argument(
        "--type",
        default="competency",
        choices=["annual", "competency"],
        help="Type of review (default: competency)"
    )
    parser.add_argument(
        "--year",
        help="Year for annual review (required for annual review type)"
    )
    parser.add_argument(
        "--format",
        default="markdown",
        choices=["markdown", "docx"],
        help="Output format (default: markdown)"
    )
    parser.add_argument(
        "--output",
        help="Custom output path for the generated report"
    )

    args = parser.parse_args()

    # Handle configuration-only commands first
    if args.create_config:
        try:
            from config_validation import ConfigValidator
            validator = ConfigValidator(args.config)
            validator._create_example_config()
            print(f"✓ Example configuration created at: {args.config}")
            print("Please edit the configuration file with your actual settings before running reviews.")
            return 0
        except Exception as e:
            print(f"Error creating configuration: {e}", file=sys.stderr)
            return 1
    
    if args.test_config:
        try:
            config = load_and_validate_config(args.config, create_if_missing=True, test_connections=True)
            print("✓ Configuration validation completed successfully!")
            return 0
        except ConfigValidationError as e:
            print(f"Configuration validation failed: {e}", file=sys.stderr)
            return 1
        except Exception as e:
            print(f"Unexpected error during configuration test: {e}", file=sys.stderr) 
            return 1

    # Validate review arguments
    if args.type == "annual" and not args.year:
        parser.error("--year is required for annual reviews")
    

    try:
        # Load and validate configuration
        print("Loading configuration...")
        try:
            config = load_and_validate_config(args.config, create_if_missing=False, test_connections=False)
        except ConfigValidationError as e:
            print(f"Configuration error: {e}", file=sys.stderr)
            print("Hint: Use --create-config to create an example configuration file.", file=sys.stderr)
            return 1
        
        print("Configuration loaded successfully.")
        
        # Validate LLM configuration
        llm_issues = validate_llm_config(config)
        if llm_issues:
            print("LLM Configuration Issues Found:")
            for issue in llm_issues:
                print(f"  - {issue}")
            print("Note: System will fall back to manual analysis if LLM fails.")
        
        # Generate the review using CSV data source
        report_path = generate_review_with_config(
            config=config,
            source="csv",
            input_file=args.file,
            review_type=args.type,
            year=args.year,
            output_format=args.format,
            output_path=args.output
        )
        print(f"\n✓ Review generation complete! The report is available at: {report_path}")
        return 0

    except ValueError as e:
        print(f"\nValidation error: {e}", file=sys.stderr)
        return 1
    except ImportError as e:
        print(f"\nDependency error: {e}", file=sys.stderr) 
        print("Hint: Ensure all required modules are installed and accessible.", file=sys.stderr)
        return 1
    except Exception as e:
        print(f"\nReview generation failed: {str(e)}", file=sys.stderr)
        print("Please check the error messages above for more details.", file=sys.stderr)
        return 1


if __name__ == "__main__":
    sys.exit(main())
